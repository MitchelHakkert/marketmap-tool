from __future__ import annotations

from collections import Counter
from datetime import datetime
from io import BytesIO
from pathlib import Path
import re
from typing import Iterable

from docx import Document
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
import streamlit as st


STANDARD_COLUMNS = [
    "naam",
    "huidig_bedrijf",
    "functietitel",
    "locatie",
    "sector",
    "tenure_huidige_rol_jaren",
    "info_eerdere_rollen",
    "hoogst_afgeronde_opleiding",
    "interesse_signalen",
    "activiteit",
]

HEADER_IMAGE_PATH = Path(__file__).with_name("assets") / "Derks-LinkedIn-achtergrond-1584x396.jpg"

DEFAULT_SECTOR_RULES = """GGZ: ggz, geestelijke gezondheidszorg, psychiatrie, psychiatrisch, verslavingszorg, autisme, mental care, parnassia, arkin, yulius, pro persona, altrecht, lentis, mondriaan, fivoor, dimence, ggze, ggz breburg, ggz centraal, tactus, ggz drenthe, ggz friesland, reinier van arkel, emergis, karakter, vnn, novadic, eleos, lievegoed, de viersprong, yes we can, hsk groep, mentaal beter
MSZ: msz, ziekenhuis, ziekenhuizen, medisch centrum, umc, universitair medisch, kliniek, klinieken, diagnostiek, laboratorium, zbc, erasmus mc, amsterdam umc, umcg, radboudumc, lumc, olvg, zuyderland, maastricht umc, canisius, st. antonius, elisabeth-tweesteden, etz, bravis, maxima mc, máxima mc, hagaziekenhuis, gelre, alrijne, adrz, tergooi, martini, bergman, diakonessenhuis, franciscus, antoni van leeuwenhoek, jeroen bosch, viecuri, catharina, rijnstate, dijklander, slingeland, zgt, flevoziekenhuis, maasstad, laurentius ziekenhuis, curaMare, curamare, dicoon, acibadem, leiden university medical, medisch spectrum twente, amphia, saxenburgh, sjg weert, caresq, rivas zorggroep, unilabs
VVT: vvt, ouderenzorg, thuiszorg, verpleeghuis, verpleeghuizen, woonzorg, wijkverpleging, zorgcentrum, zorgcentra, brabantzorg, careyn, laurens, cordaan, aafje, tantelouise, de zorgcirkel, beweging 3.0, zonnehuisgroep, zorgsaam, actief zorg, bloezem, lelie zorggroep, surplus, mijzo, zorggroep solis, carinova, zinn, magentazorg, florence, coloriet, sevagram, van neynsel, domus valuas, groenhuysen, pantein, pieter van foreest, welthuis, amstelring, dignis, axioncontinu, sensire, meandergroep, amarijn, saffier, kwadrantgroep, omring, cicero, humanitas, opella, viva! zorggroep, amsta, hilverzorg, avoord, topaz, zorggroep ter weel, korian, marente, activite, zuidoostzorg, zorg groep beek, miep huishoudelijke, zorgspectrum, vilente, viattence, bartholomeus gasthuis, zorgspectrum het zand, quarijn, stichting land van horne
WLZ: wlz, gehandicaptenzorg, beperking, beperkingen, lvb, langdurige zorg, forensische zorg, jeugdwet, jeugdzorg, pluryn, middin, cosis, 's heeren loo, cello, siza, amarant, gemiva, ipse de bruggen, koninklijke visio, ambiq, timon, ons tweede thuis, triade vitree, zuidwester, stichting radar, wender, sdw zorg, oro, lunet, zozijn, omega groep, tragel, prinsenstichting, sovak, sensa zorg, zideris, aveleijn, driestroom, baalderborg, abrona, philadelphia, ribw, raphaëlstichting, raphaelstichting
Pharma: pharma, pharmaceutical, geneesmiddel, geneesmiddelen, farma, fabrikant, manufacturing, batch release, gmp, qa, quality assurance, qualified person, qp
Biotech: biotech, biotechnology, biologics, biologisch, cell therapy, gene therapy
CDMO: cdmo, cmo, contract manufacturing, contract development, fill finish
Medical Devices: medical device, medical devices, medtech, hulpmiddel, hulpmiddelen, device
Healthcare: healthcare, health care, zorg, ziekenhuis, clinic, kliniek
Overig: overig"""

SECTOR_PRESETS = {
    "Automatisch": {
        "buckets": [],
        "include_rule_buckets": [],
    },
    "Zorg": {
        "buckets": ["GGZ", "MSZ", "VVT", "WLZ", "Overig"],
        "include_rule_buckets": ["GGZ", "MSZ", "VVT", "WLZ", "Overig"],
    },
    "Pharma / biotech": {
        "buckets": ["Pharma", "Biotech", "CDMO", "Medical Devices", "Healthcare", "Overig"],
        "include_rule_buckets": ["Pharma", "Biotech", "CDMO", "Medical Devices", "Healthcare", "Overig"],
    },
}

DEFAULT_INCLUDE_TERMS = {
    "Zorg": [
        "zorgverkoop",
        "zorgverkoper",
        "zorgcontractering",
        "zorgcontract",
        "contractering",
        "zorginkoop",
        "zorginkoper",
        "zorgbemiddeling",
        "zorgbemiddelaar",
        "accountmanager zorg",
        "adviseur zorgverkoop",
        "strategisch zorgverkoper",
        "relatiebeheer",
        "wlz",
        "zvw",
        "wmo",
        "jeugdwet",
    ],
    "Pharma / biotech": [
        "qualified person",
        "qp",
        "q.p.",
        "gmp",
        "batch release",
        "qa",
        "quality assurance",
        "quality",
        "regulatory affairs",
        "regulatory",
        "pharma",
        "pharmaceutical",
        "biotech",
        "cdmo",
        "medical device",
        "medtech",
    ],
}

DEFAULT_EXCLUDE_TERMS = [
    "recruiter",
    "student",
    "stagiair",
    "intern",
    "finance consultant",
    "treasury",
    "opticien",
    "ergotherapeut",
    "fysiotherapeut",
    "verpleegkundige",
    "docent",
    "hogeschooldocent",
    "raad van toezicht",
]

AUDIENCE_PATTERNS = [
    ("Zorgverkopers", ["zorgverkoop", "zorgverkoper", "adviseur zorgverkoop", "accountmanager zorgverkoop"]),
    ("Zorginkopers", ["zorginkoop", "zorginkoper", "inkoper zorg", "zorgcontractering"]),
    ("Medical Affairs", ["medical affairs", "medical science liaison", "msl", "medical advisor", "medical manager"]),
    ("Qualified Persons", ["qualified person", "qp", "q.p.", "batch release"]),
    ("Regulatory Affairs", ["regulatory affairs", "regulatory", "registratie"]),
    ("Quality Assurance", ["quality assurance", "qa", "gmp", "quality manager"]),
    ("Market Access", ["market access", "health economics", "heor", "reimbursement", "pricing"]),
    ("Accountmanagers", ["accountmanager", "key account", "sales manager", "business development"]),
    ("Inkopen", ["procurement", "inkoop", "inkoper", "buyer", "sourcing"]),
]


BADGES = (
    "Dit profiel heeft een LinkedIn Premium-abonnement.",
    "Derdegraads connectie",
    "Tweedegraads connectie",
    "Eerstegraads connectie",
    "Gesourced",
    "This profile has a LinkedIn Premium subscription",
    "Third degree connection",
    "Second degree connection",
    "First degree connection",
    "Sourced",
)

STOP_PREFIXES = (
    "Stadium wijzigen",
    "Bericht sturen naar ",
    "Meer acties voor ",
    "Inuncontacted",
    "Opgeslagen op ",
    "door ",
    "Change stage",
    "Message ",
    "More actions for ",
    "Saved on ",
    "by ",
)


def clean(text: object) -> str:
    text = "" if text is None else str(text)
    text = text.replace("\xa0", " ").replace("Â·", "·").replace("â€“", "–")
    return " ".join(text.split()).strip()


def tokens(text: str) -> list[str]:
    return [clean(x).lower() for x in re.split(r"[,;\n]+", text or "") if clean(x)]


def parse_sector_rules(text: str, allowed_buckets: list[str]) -> dict[str, list[str]]:
    allowed = {bucket.lower(): bucket for bucket in allowed_buckets}
    rules: dict[str, list[str]] = {bucket: [] for bucket in allowed_buckets}
    for raw_line in (text or "").splitlines():
        line = clean(raw_line)
        if not line or ":" not in line:
            continue
        bucket_raw, keywords_raw = line.split(":", 1)
        bucket = allowed.get(clean(bucket_raw).lower())
        if not bucket:
            continue
        rules[bucket].extend(tokens(keywords_raw))
    return rules


def sector_preset_for_df(df: pd.DataFrame) -> str:
    blob = " ".join(
        " ".join(clean(row.get(c)) for c in ["huidig_bedrijf", "functietitel", "sector", "info_eerdere_rollen"])
        for _, row in df.head(200).iterrows()
    ).lower()
    zorg_score = sum(
        blob.count(term)
        for term in ["zorgverkoop", "zorgcontract", "ggz", "ziekenhuis", "ouderenzorg", "wlz", "vvt", "parnassia", "pluryn"]
    )
    pharma_score = sum(
        blob.count(term)
        for term in ["qualified person", "batch release", "gmp", "pharma", "pharmaceutical", "biotech", "cdmo", "medical device"]
    )
    return "Pharma / biotech" if pharma_score > zorg_score else "Zorg"


def sector_config(preset_name: str, df: pd.DataFrame) -> tuple[list[str], dict[str, list[str]], str]:
    resolved = sector_preset_for_df(df) if preset_name == "Automatisch" else preset_name
    preset = SECTOR_PRESETS[resolved]
    buckets = preset["buckets"]
    rules = parse_sector_rules(DEFAULT_SECTOR_RULES, preset["include_rule_buckets"])
    return buckets, rules, resolved


def is_badge(text: str) -> bool:
    return any(text.startswith(prefix) for prefix in BADGES)


def is_sector_line(text: str) -> bool:
    return text.startswith(("Â· ", "· "))


def clean_sector_line(text: str) -> str:
    return clean(re.sub(r"^(?:Â·|·)\s*", "", text))


def is_experience_marker(text: str) -> bool:
    return text in {"ErvaringErvaring in profiel", "ExperienceProfile experience"}


def is_education_marker(text: str) -> bool:
    return text in {"OpleidingOpleidingen op het profiel", "EducationProfile education"}


def is_interest_marker(text: str) -> bool:
    return text.startswith(
        (
            "LabelsLabels",
            "Decoraties voor rij met profielinteresses",
            "TagsCandidate tags",
            "Profile interest row decorations",
        )
    )


def is_activity_marker(text: str) -> bool:
    return text.startswith(("Decoraties voor rij met profielactiviteit", "Profile activity row decorations"))



def strip_date_suffix(line: str) -> str:
    return clean(re.sub(r"\s*(?:Â·|·|愧)\s*.*$", "", line))


def split_title_company(headline: str) -> tuple[str, str]:
    h = clean(headline)
    for sep in (" bij ", " at ", " @ "):
        if sep in h:
            title, company = h.rsplit(sep, 1)
            return clean(title), clean(company)
    return h, ""


def years_from_line(line: str) -> int | None:
    if not line:
        return None
    years = [int(y) for y in re.findall(r"\b(19\d{2}|20\d{2})\b", line)]
    if not years:
        return None
    start = years[0]
    end = datetime.now().year if re.search(r"\bheden\b|\bpresent\b", line, re.I) else years[-1]
    return max(0, end - start)


HBO_INSTITUTIONS = (
    "hogeschool",
    "university of applied sciences",
    "han ",
    "hogeschool van arnhem",
    "arnhem en nijmegen",
    "hanze",
    "fontys",
    "avans",
    "saxion",
    "inholland",
    "windesheim",
    "nhl stenden",
    "stenden",
    "hva",
    "hu university",
    "hogeschool utrecht",
    "hogeschool rotterdam",
    "rotterdam university of applied",
    "the hague university of applied",
    "amsterdam university of applied",
    "zuyd",
    "hz university",
    "has green academy",
    "aeres",
    "breda university of applied",
    "nhtv",
    "artez",
    "codarts",
    "hotelschool the hague",
    "tio university",
    "wittenborg",
    "viaa",
    "che ",
    "christelijke hogeschool ede",
    "ipabo",
    "marnix academie",
    "driestar",
    "iselinge",
    "kempel",
    "gerrit rietveld academie",
    "design academy eindhoven",
    "amsterdamse hogeschool voor de kunsten",
    "ncoi",
    "isbw",
    "loi",
    "nti",
    "scheidegger",
)

WO_INSTITUTIONS = (
    "universiteit",
    "university",
    "erasmus",
    "nyenrode",
    "tilburg university",
    "utrecht university",
    "university of amsterdam",
    "vrije universiteit",
    "vu amsterdam",
    "leiden university",
    "rijksuniversiteit groningen",
    "university of groningen",
    "radboud",
    "maastricht university",
    "university of twente",
    "delft university",
    "tu delft",
    "wageningen university",
    "eindhoven university",
    "open universiteit",
    "rijksuniversiteit",
    "universiteit leiden",
    "universiteit utrecht",
    "universiteit van amsterdam",
    "universiteit twente",
    "technische universiteit",
    "tu eindhoven",
    "tilburg universiteit",
    "maastricht universiteit",
    "radboud universiteit",
    "vrije universiteit amsterdam",
    "wageningen universiteit",
)

MBO_INSTITUTIONS = (
    "roc ",
    "roc-",
    "regionaal opleidingencentrum",
    "alfa-college",
    "alfa college",
    "zadkine",
    "menso alting",
    "koning willem i college",
    "horizon college",
    "roc horizon",
    "roc a12",
    "mondriaan",
    "rocmondriaan",
    "mborijnland",
    "mbo rijnland",
    "nova college",
    "deltion",
    "summa college",
    "vista college",
    "gilde opleidingen",
    "friesland college",
    "firda",
    "curio",
    "noorderpoort",
    "drenthe college",
    "terra mbo",
    "graafschap college",
    "aventus",
    "rijn ijssel",
    "scalda",
    "da vinci college",
    "talland college",
    "yuverta",
    "sintlucas",
    "nimeto",
    "cibap",
    "vakschool",
    "de rooi pannen",
    "middelbare hotelschool",
    "hotelschool groningen",
)

EDUCATION_ORDER = ["PhD/PostDoc", "WO master", "WO Bachelor", "HBO", "MBO", "Overig"]


def infer_education_level(line: str) -> str:
    t = f" {clean(line).lower()} "
    is_hbo_school = any(school in t for school in HBO_INSTITUTIONS)
    is_wo_school = any(school in t for school in WO_INSTITUTIONS) and not is_hbo_school
    is_mbo_school = any(school in t for school in MBO_INSTITUTIONS)

    if re.search(r"\b(phd|ph\.d|doctorate|postdoc|post-doc|postdoctoral|promotie|gepromoveerd)\b", t):
        return "PhD/PostDoc"
    if (
        is_mbo_school
        or re.search(r"\b(mbo|middelbaar beroepsonderwijs|niveau\s*[1-4]|level\s*[1-4]|mbo[- ]?niveau)\b", t)
        or re.search(r"\b(mdgo|meao|mavo|vmbo|helpende|verzorgende\s+ig|verpleegkundige\s+niveau\s*[1-4])\b", t)
    ):
        return "MBO"
    if is_hbo_school or re.search(r"\b(hbo|heao|hts|hbs)\b", t):
        return "HBO"
    if re.search(r"\b(maatschappelijk werk en dienstverlening|social work|sociaal pedagogische hulpverlening|sph|mwd)\b", t):
        return "HBO"
    if re.search(r"\b(master|msc|m\.sc|ms|ma|m\.a\.|llm|m\.?b\.?a\.?|drs|doctoraal)\b", t):
        return "WO master"
    if re.search(r"\b(bsc|b\.sc|ba|b\.a\.)\b", t):
        return "WO Bachelor" if is_wo_school else "HBO"
    if re.search(r"\b(bachelor|bachlor|bachelorgraad|bba|post bachelor|post-hbo|associate degree|associate's degree|ad)\b", t):
        return "WO Bachelor" if is_wo_school and not is_hbo_school else "HBO"
    if is_wo_school or re.search(r"\b(wo|doctoraal)\b", t):
        return "WO master"
    return ""


def education_summary(lines: list[str]) -> str:
    if not lines:
        return "Overig"
    rank = {"PhD/PostDoc": 5, "WO master": 4, "WO Bachelor": 3, "HBO": 2, "MBO": 1}
    levels = [infer_education_level(line) for line in lines]
    known_levels = [level for level in levels if level]
    if not known_levels:
        return "Overig"
    return max(known_levels, key=lambda level: rank.get(level, 0))


def normalize_education_value(value: object) -> str:
    text = clean(value)
    if not text or text.lower() in {"onbekend", "unknown", "overig", "other"}:
        return "Overig"
    for level in EDUCATION_ORDER:
        if text.lower() == level.lower():
            return level
    return infer_education_level(text) or "Overig"


def compact_role(line: str) -> str:
    years = re.findall(r"\b(19\d{2}|20\d{2})\b", line or "")
    role = strip_date_suffix(line)
    if len(years) >= 2:
        return f"{role} ({years[0]}-{years[-1]})"
    if len(years) == 1:
        return f"{role} ({years[0]})"
    return role


def best_current_experience(experience: list[str], title: str, company: str) -> str:
    current = [x for x in experience if re.search(r"\bheden\b|\bpresent\b", x, re.I)]
    search = current or experience
    title_l, company_l = title.lower(), company.lower()
    for line in search:
        ll = line.lower()
        if title_l and title_l in ll:
            return line
    for line in search:
        ll = line.lower()
        if company_l and company_l in ll:
            return line
    return search[0] if search else ""


def _legacy_parse_docx_unused(uploaded_file) -> pd.DataFrame:
    doc = Document(uploaded_file)
    texts = [clean(p.text) for p in doc.paragraphs]
    texts = [t for t in texts if t]
    exp_idxs = [i for i, t in enumerate(texts) if t == "ErvaringErvaring in profiel"]

    records = []
    if not exp_idxs:
        return pd.DataFrame(columns=STANDARD_COLUMNS)

    headers = []
    for exp_idx in exp_idxs:
        if exp_idx < 4:
            continue
        sector = texts[exp_idx - 1][2:].strip() if texts[exp_idx - 1].startswith("· ") else ""
        location = texts[exp_idx - 2]
        headline = texts[exp_idx - 3]
        name_idx = exp_idx - 4
        while name_idx >= 0 and is_badge(texts[name_idx]):
            name_idx -= 1
        if texts[name_idx].endswith(" selecteren") and name_idx + 1 < len(texts):
            name_idx += 1
        headers.append((name_idx, texts[name_idx], headline, location, sector))

    for pos, (start, name, headline, location, sector) in enumerate(headers):
        end = headers[pos + 1][0] if pos + 1 < len(headers) else len(texts)
        block = texts[start:end]
        info = [x for x in block[1:] if not is_badge(x)]
        experience: list[str] = []
        education: list[str] = []
        interest: list[str] = []
        activity: list[str] = []
        section = None

        for line in info:
            if line.startswith("· "):
                continue
            if section is None and line in {headline, location}:
                continue
            if line == "ErvaringErvaring in profiel":
                section = "experience"
                continue
            if line == "OpleidingOpleidingen op het profiel":
                section = "education"
                continue
            if line.startswith("LabelsLabels") or line.startswith("Decoraties voor rij met profielinteresses"):
                section = "interest"
                continue
            if line.startswith("Decoraties voor rij met profielactiviteit"):
                section = "activity"
                continue
            if line.startswith("Alles weergeven"):
                continue
            if any(line.startswith(prefix) for prefix in STOP_PREFIXES):
                section = None
                continue
            if line.startswith("Kandidaat werd opgeslagen") or line.startswith("Opgeslagen door"):
                section = "activity"
                continue
            if section == "experience":
                experience.append(line)
            elif section == "education":
                education.append(line)
            elif section == "interest":
                interest.append(line)
            elif section == "activity":
                activity.append(line)

        title, company = split_title_company(headline)
        current = best_current_experience(experience, title, company)
        if current:
            current_title, current_company = split_title_company(strip_date_suffix(current))
            if current_company and not company:
                company = current_company
            if current_company and title.lower() == current_company.lower():
                title = current_title

        prior = [x for x in experience if x != current]
        if "Premium" in " ".join(block):
            interest.append("LinkedIn Premium")

        records.append(
            {
                "naam": name,
                "huidig_bedrijf": company,
                "functietitel": title,
                "locatie": location,
                "sector": sector,
                "tenure_huidige_rol_jaren": years_from_line(current),
                "info_eerdere_rollen": " | ".join(compact_role(x) for x in prior[:3]),
                "hoogst_afgeronde_opleiding": education_summary(education),
                "interesse_signalen": "; ".join(dict.fromkeys([x for x in interest if x])),
                "activiteit": " | ".join(activity[:4]),
            }
        )
    return pd.DataFrame(records, columns=STANDARD_COLUMNS)


def parse_docx(uploaded_file) -> pd.DataFrame:
    doc = Document(uploaded_file)
    texts = [clean(p.text) for p in doc.paragraphs]
    texts = [t for t in texts if t]
    exp_idxs = [i for i, t in enumerate(texts) if is_experience_marker(t)]

    records = []
    if not exp_idxs:
        return pd.DataFrame(columns=STANDARD_COLUMNS)

    headers = []
    for exp_idx in exp_idxs:
        if exp_idx < 3:
            continue

        scan_idx = exp_idx - 1
        sector = ""
        if scan_idx >= 0 and is_sector_line(texts[scan_idx]):
            sector = clean_sector_line(texts[scan_idx])
            scan_idx -= 1
        if scan_idx < 2:
            continue

        location = texts[scan_idx]
        headline = texts[scan_idx - 1]
        name_idx = scan_idx - 2
        while name_idx >= 0 and (is_badge(texts[name_idx]) or texts[name_idx].startswith("Select ")):
            name_idx -= 1
        if name_idx < 0:
            continue
        if texts[name_idx].endswith(" selecteren") and name_idx + 1 < len(texts):
            name_idx += 1
        headers.append((name_idx, texts[name_idx], headline, location, sector))

    for pos, (start, name, headline, location, sector) in enumerate(headers):
        end = headers[pos + 1][0] if pos + 1 < len(headers) else len(texts)
        block = texts[start:end]
        info = [x for x in block[1:] if not is_badge(x)]
        experience: list[str] = []
        education: list[str] = []
        interest: list[str] = []
        activity: list[str] = []
        section = None

        for line in info:
            if is_sector_line(line):
                continue
            if section is None and line in {headline, location}:
                continue
            if is_experience_marker(line):
                section = "experience"
                continue
            if is_education_marker(line):
                section = "education"
                continue
            if is_interest_marker(line):
                section = "interest"
                continue
            if is_activity_marker(line):
                section = "activity"
                continue
            if line.startswith(("Alles weergeven", "Show all", "Similar skills to saved candidates")):
                continue
            if any(line.startswith(prefix) for prefix in STOP_PREFIXES):
                section = None
                continue
            if line.startswith(("Kandidaat werd opgeslagen", "Opgeslagen door", "Candidate saved to project", "Saved by")):
                section = "activity"
                continue
            if section == "experience":
                experience.append(line)
            elif section == "education":
                education.append(line)
            elif section == "interest":
                interest.append(line)
            elif section == "activity":
                activity.append(line)

        title, company = split_title_company(headline)
        current = best_current_experience(experience, title, company)
        if current:
            current_title, current_company = split_title_company(strip_date_suffix(current))
            if current_company and not company:
                company = current_company
            if current_company and title.lower() == current_company.lower():
                title = current_title

        prior = [x for x in experience if x != current]
        if "Premium" in " ".join(block):
            interest.append("LinkedIn Premium")

        records.append(
            {
                "naam": name,
                "huidig_bedrijf": company,
                "functietitel": title,
                "locatie": location,
                "sector": sector,
                "tenure_huidige_rol_jaren": years_from_line(current),
                "info_eerdere_rollen": " | ".join(compact_role(x) for x in prior[:3]),
                "hoogst_afgeronde_opleiding": education_summary(education),
                "interesse_signalen": "; ".join(dict.fromkeys([x for x in interest if x])),
                "activiteit": " | ".join(activity[:4]),
            }
        )
    return pd.DataFrame(records, columns=STANDARD_COLUMNS)


def read_table(uploaded_file) -> pd.DataFrame:
    name = uploaded_file.name.lower()
    if name.endswith(".docx"):
        df = parse_docx(uploaded_file)
    elif name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
        df = normalize_columns(df)
    else:
        df = pd.read_excel(uploaded_file)
        df = normalize_columns(df)
    df["hoogst_afgeronde_opleiding"] = df["hoogst_afgeronde_opleiding"].apply(normalize_education_value)
    return df[STANDARD_COLUMNS]


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    aliases = {
        "naam": ["naam", "name", "candidate", "kandidaat"],
        "huidig_bedrijf": ["huidig bedrijf", "company", "bedrijf", "current company", "organisatie"],
        "functietitel": ["functietitel", "title", "functie", "job title", "headline"],
        "locatie": ["locatie", "location", "plaats"],
        "sector": ["sector", "industry", "branche"],
        "tenure_huidige_rol_jaren": ["tenure huidige rol (jaren)", "tenure", "tenure_huidige_rol_jaren"],
        "info_eerdere_rollen": ["info eerdere rollen", "previous roles", "ervaring"],
        "hoogst_afgeronde_opleiding": ["hoogst afgeronde opleiding", "education", "opleiding"],
        "interesse_signalen": ["interesse signalen", "interest", "signals"],
        "activiteit": ["activiteit", "activity"],
    }
    by_lower = {clean(c).lower(): c for c in df.columns}
    out = pd.DataFrame()
    for target, names in aliases.items():
        source = next((by_lower[n] for n in names if n in by_lower), None)
        out[target] = df[source] if source else ""
    return out[STANDARD_COLUMNS]


def classify_sector(row: pd.Series, sector_buckets: list[str], sector_rules: dict[str, list[str]] | None = None) -> str:
    if not sector_buckets:
        return clean(row.get("sector")) or "Overig"
    blob = " | ".join(clean(row.get(c)) for c in ["huidig_bedrijf", "functietitel", "info_eerdere_rollen"]).lower()
    original_sector = clean(row.get("sector")).lower()
    for bucket in sector_buckets:
        keywords = (sector_rules or {}).get(bucket, [])
        if any(keyword and keyword in blob for keyword in keywords):
            return bucket
    for bucket in sector_buckets:
        bucket_l = bucket.lower()
        bucket_words = [bucket_l, *bucket_l.replace("/", " ").split()]
        if any(word and word in blob for word in bucket_words):
            return bucket
        if original_sector == bucket_l:
            return bucket
    return "Overig" if "Overig" in sector_buckets else sector_buckets[-1]


OUTLIER_STOPWORDS = {
    "account",
    "accountmanager",
    "adviseur",
    "alleen",
    "assistant",
    "bedrijf",
    "bij",
    "business",
    "consultant",
    "coordinator",
    "current",
    "department",
    "director",
    "ervaring",
    "functie",
    "health",
    "healthcare",
    "huidig",
    "manager",
    "medewerker",
    "nederland",
    "netherlands",
    "opleiding",
    "project",
    "senior",
    "specialist",
    "team",
    "van",
    "voor",
    "werk",
}


def learned_relevance_terms(df: pd.DataFrame, minimum_share: float = 0.08) -> list[str]:
    texts = []
    for _, row in df.iterrows():
        texts.append(
            " ".join(
                clean(row.get(col))
                for col in ["huidig_bedrijf", "functietitel", "sector", "info_eerdere_rollen"]
            ).lower()
        )
    min_count = max(3, int(len(texts) * minimum_share))
    counts: Counter[str] = Counter()
    for text in texts:
        words = {
            word
            for word in re.findall(r"[a-zA-ZÀ-ÿ][a-zA-ZÀ-ÿ0-9\-]{3,}", text)
            if word not in OUTLIER_STOPWORDS
        }
        counts.update(words)
    return [word for word, count in counts.most_common(30) if count >= min_count]


def term_hits(terms: list[str], text: str) -> list[str]:
    hits = []
    for term in terms:
        clean_term = clean(term).lower()
        if not clean_term:
            continue
        pattern = r"(?<![a-zA-ZÀ-ÿ0-9])" + re.escape(clean_term) + r"(?![a-zA-ZÀ-ÿ0-9])"
        if re.search(pattern, text):
            hits.append(clean_term)
    return hits


def guess_audience_name(df: pd.DataFrame, fallback: str = "Profielen") -> str:
    blob = " ".join(
        clean(row.get(col))
        for _, row in df.iterrows()
        for col in ["huidig_bedrijf", "functietitel", "sector", "info_eerdere_rollen", "interesse_signalen"]
    ).lower()
    scores = []
    for label, terms in AUDIENCE_PATTERNS:
        score = sum(len(term_hits([term], blob)) for term in terms)
        if score:
            scores.append((score, label))
    if scores:
        return max(scores)[1]

    learned = learned_relevance_terms(df, minimum_share=0.12)
    if learned:
        words = [word.capitalize() for word in learned[:2]]
        return " ".join(words)
    return fallback


def safe_excel_filename(label: str) -> str:
    label = clean(label) or "Profielen"
    label = re.sub(r'[<>:"/\\|?*]+', "", label)
    label = re.sub(r"\s+", " ", label).strip(" ._-")
    return f"Market Map {label or 'Profielen'}.xlsx"


def safe_pdf_filename(label: str) -> str:
    label = clean(label) or "Profielen"
    label = re.sub(r'[<>:"/\\|?*]+', "", label)
    label = re.sub(r"\s+", " ", label).strip(" ._-")
    return f"Market Map {label or 'Profielen'} - Onepager.pdf"


def make_onepager_pdf(df: pd.DataFrame, audience_name: str = "") -> bytes:
    output = BytesIO()
    page_width, page_height = landscape(A4)
    pdf = canvas.Canvas(output, pagesize=landscape(A4))

    dark_blue = colors.HexColor("#073763")
    mid_blue = colors.HexColor("#155E75")
    light_blue = colors.HexColor("#EAF4FA")
    green = colors.HexColor("#92D050")
    grey = colors.HexColor("#4B5563")
    border = colors.HexColor("#D9E2EA")

    margin = 28
    header_height = 142
    if HEADER_IMAGE_PATH.exists():
        image = ImageReader(str(HEADER_IMAGE_PATH))
        pdf.drawImage(image, 0, page_height - header_height, width=page_width, height=header_height, preserveAspectRatio=True, anchor="n")
    else:
        pdf.setFillColor(dark_blue)
        pdf.rect(0, page_height - header_height, page_width, header_height, fill=1, stroke=0)
        pdf.setFillColor(colors.white)
        pdf.setFont("Helvetica-Bold", 24)
        pdf.drawCentredString(page_width / 2, page_height - 75, "Derks & Derks")
    pdf.setFillColor(green)
    pdf.rect(0, page_height - header_height - 10, page_width, 10, fill=1, stroke=0)

    title_y = page_height - header_height - 42
    title = f"Market Map {clean(audience_name) or 'Profielen'}"
    pdf.setFillColor(dark_blue)
    pdf.setFont("Helvetica-Bold", 22)
    pdf.drawString(margin, title_y, title)
    pdf.setFillColor(grey)
    pdf.setFont("Helvetica", 9)
    pdf.drawString(margin, title_y - 15, "Kernbeeld van de profielen in deze mapping")

    total = max(len(df), 1)
    tenure = pd.to_numeric(df["tenure_huidige_rol_jaren"], errors="coerce")
    avg_tenure = tenure.mean()
    top_company = df["huidig_bedrijf"].replace("", "Onbekend").fillna("Onbekend").value_counts().index[0] if len(df) else "-"
    top_sector = df["sector"].replace("", "Overig").fillna("Overig").value_counts().index[0] if len(df) else "-"

    def draw_card(x: float, y: float, w: float, h: float, label: str, value: str) -> None:
        pdf.setFillColor(mid_blue)
        pdf.roundRect(x, y + h - 22, w, 22, 3, fill=1, stroke=0)
        pdf.setFillColor(colors.white)
        pdf.setFont("Helvetica-Bold", 8)
        pdf.drawCentredString(x + w / 2, y + h - 15, label)
        pdf.setFillColor(light_blue)
        pdf.roundRect(x, y, w, h - 22, 3, fill=1, stroke=1)
        pdf.setStrokeColor(border)
        pdf.setFillColor(dark_blue)
        pdf.setFont("Helvetica-Bold", 15)
        clipped = value if len(value) <= 22 else value[:21] + "..."
        pdf.drawCentredString(x + w / 2, y + 18, clipped)

    card_y = title_y - 72
    card_w = 120
    gap = 14
    draw_card(margin, card_y, card_w, 56, "Profielen", str(len(df)))
    draw_card(margin + (card_w + gap), card_y, card_w, 56, "Gem. tenure", "" if pd.isna(avg_tenure) else f"{avg_tenure:.1f} jaar")
    draw_card(margin + 2 * (card_w + gap), card_y, card_w, 56, "Topbedrijf", str(top_company))
    draw_card(margin + 3 * (card_w + gap), card_y, card_w, 56, "Topsector", str(top_sector))

    def draw_table(x: float, y: float, title_text: str, rows: list[tuple[str, int, float]], col_widths: tuple[int, int, int]) -> None:
        row_h = 17
        table_w = sum(col_widths)
        pdf.setFillColor(dark_blue)
        pdf.rect(x, y, table_w, row_h, fill=1, stroke=0)
        pdf.setFillColor(colors.white)
        pdf.setFont("Helvetica-Bold", 9)
        pdf.drawString(x + 5, y + 5, title_text)
        header_y = y - row_h
        pdf.setFillColor(colors.HexColor("#DDEBF7"))
        pdf.rect(x, header_y, table_w, row_h, fill=1, stroke=1)
        pdf.setFillColor(colors.black)
        pdf.setFont("Helvetica-Bold", 7)
        headers = ["categorie", "aantal", "%"]
        cx = x
        for header, width in zip(headers, col_widths):
            pdf.drawString(cx + 4, header_y + 5, header)
            cx += width
        pdf.setFont("Helvetica", 7)
        for idx, (label, count, pct) in enumerate(rows):
            row_y = header_y - row_h * (idx + 1)
            pdf.setFillColor(colors.white)
            pdf.rect(x, row_y, table_w, row_h, fill=1, stroke=1)
            pdf.setFillColor(colors.black)
            clipped = str(label) if len(str(label)) <= 28 else str(label)[:27] + "..."
            pdf.drawString(x + 4, row_y + 5, clipped)
            pdf.drawRightString(x + col_widths[0] + col_widths[1] - 5, row_y + 5, str(count))
            pdf.drawRightString(x + table_w - 5, row_y + 5, f"{pct:.1%}")

    edu_counts = df["hoogst_afgeronde_opleiding"].replace("", "Overig").fillna("Overig").value_counts()
    education_rows = [(label, int(edu_counts.get(label, 0)), int(edu_counts.get(label, 0)) / total) for label in EDUCATION_ORDER]
    company_counts = df["huidig_bedrijf"].replace("", "Onbekend").fillna("Onbekend").value_counts().head(5)
    company_rows = [(label, int(count), int(count) / total) for label, count in company_counts.items()]
    sector_counts = df["sector"].replace("", "Overig").fillna("Overig").value_counts().head(5)
    sector_rows = [(label, int(count), int(count) / total) for label, count in sector_counts.items()]
    tenure_rows = [
        ("<1 jaar", int(((tenure >= 0) & (tenure < 1)).sum()), int(((tenure >= 0) & (tenure < 1)).sum()) / total),
        ("1-2 jaar", int(((tenure >= 1) & (tenure <= 2)).sum()), int(((tenure >= 1) & (tenure <= 2)).sum()) / total),
        ("3-5 jaar", int(((tenure >= 3) & (tenure <= 5)).sum()), int(((tenure >= 3) & (tenure <= 5)).sum()) / total),
        ("6-10 jaar", int(((tenure >= 6) & (tenure <= 10)).sum()), int(((tenure >= 6) & (tenure <= 10)).sum()) / total),
        ("10+ jaar", int((tenure > 10).sum()), int((tenure > 10).sum()) / total),
        ("Onbekend", int(tenure.isna().sum()), int(tenure.isna().sum()) / total),
    ]

    table_y = card_y - 36
    draw_table(margin, table_y, "Opleiding", education_rows, (112, 36, 42))
    draw_table(margin + 220, table_y, "Top bedrijven", company_rows, (126, 36, 42))
    draw_table(margin, table_y - 160, "Tenure", tenure_rows, (112, 36, 42))
    draw_table(margin + 220, table_y - 160, "Top sectoren", sector_rows, (126, 36, 42))

    note_x = margin + 460
    note_y = table_y
    note_w = page_width - note_x - margin
    note_h = 310
    pdf.setFillColor(dark_blue)
    pdf.rect(note_x, note_y, note_w, 18, fill=1, stroke=0)
    pdf.setFillColor(colors.white)
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(note_x + 5, note_y + 5, "Duiding / klantnotitie")
    pdf.setFillColor(colors.HexColor("#F7FBF2"))
    pdf.setStrokeColor(green)
    pdf.rect(note_x, note_y - note_h, note_w, note_h, fill=1, stroke=1)
    pdf.setFillColor(grey)
    pdf.setFont("Helvetica", 8)
    pdf.drawString(note_x + 9, note_y - 18, "Vul hier zelf de belangrijkste interpretatie, nuance of klantboodschap in.")

    pdf.showPage()
    pdf.save()
    return output.getvalue()


def flag_outliers(df: pd.DataFrame, include_terms: list[str], exclude_terms: list[str]) -> pd.DataFrame:
    learned_terms = learned_relevance_terms(df)
    rows = []
    for idx, row in df.iterrows():
        blob = " | ".join(clean(row.get(c)) for c in STANDARD_COLUMNS).lower()
        relevance_blob = " | ".join(
            clean(row.get(c))
            for c in ["huidig_bedrijf", "functietitel", "sector", "info_eerdere_rollen", "interesse_signalen"]
        ).lower()
        exclude_blob = clean(row.get("functietitel")).lower()
        include_hits = term_hits(include_terms, relevance_blob)
        learned_hits = term_hits(learned_terms, relevance_blob)
        exclude_hits = term_hits(exclude_terms, exclude_blob)
        has_sector_signal = clean(row.get("sector")).lower() not in {"", "overig", "unknown", "onbekend"}
        reasons = []
        score = 0
        if exclude_hits:
            score += 2
            reasons.append("uitsluitingskeyword(s): " + ", ".join(exclude_hits[:5]))
        if not clean(row.get("huidig_bedrijf")) or not clean(row.get("functietitel")):
            score += 2
            reasons.append("bedrijf of functietitel ontbreekt")
        if include_terms and not include_hits and not learned_hits and not has_sector_signal:
            score += 2
            reasons.append("geen standaard- of input-relevant woord gevonden")
        elif include_terms and not include_hits and not has_sector_signal:
            score += 1
            reasons.append("geen standaarddoelwoord, maar wel inputsignaal: " + ", ".join(learned_hits[:5]))
        if score >= 2:
            rows.append(
                {
                    "verwijderen": False,
                    "index": idx,
                    "twijfelscore": score,
                    "naam": row.get("naam", ""),
                    "huidig_bedrijf": row.get("huidig_bedrijf", ""),
                    "functietitel": row.get("functietitel", ""),
                    "sector": row.get("sector", ""),
                    "reden": "; ".join(reasons),
                    "locatie": row.get("locatie", ""),
                    "tenure_huidige_rol_jaren": row.get("tenure_huidige_rol_jaren", ""),
                    "info_eerdere_rollen": row.get("info_eerdere_rollen", ""),
                    "hoogst_afgeronde_opleiding": row.get("hoogst_afgeronde_opleiding", ""),
                    "interesse_signalen": row.get("interesse_signalen", ""),
                    "activiteit": row.get("activiteit", ""),
                }
            )
    return pd.DataFrame(rows)


def summary_frames(df: pd.DataFrame) -> dict[str, pd.DataFrame]:
    total = max(len(df), 1)
    frames = {
        "KPI": pd.DataFrame(
            [
                ["Profielen", len(df), 1],
                ["Met huidig bedrijf", df["huidig_bedrijf"].astype(str).str.len().gt(0).sum(), None],
                ["Met tenure", pd.to_numeric(df["tenure_huidige_rol_jaren"], errors="coerce").notna().sum(), None],
                ["Gemiddelde tenure", pd.to_numeric(df["tenure_huidige_rol_jaren"], errors="coerce").mean(), None],
                ["Met opleiding", df["hoogst_afgeronde_opleiding"].astype(str).str.len().gt(0).sum(), None],
            ],
            columns=["metriek", "waarde", "percentage"],
        )
    }
    for col, label in [("huidig_bedrijf", "Bedrijven"), ("sector", "Sectoren"), ("locatie", "Locaties")]:
        counts = df[col].replace("", "Onbekend").fillna("Onbekend").value_counts().head(20)
        frames[label] = pd.DataFrame({"categorie": counts.index, "aantal": counts.values, "percentage": counts.values / total})
    return frames


def make_excel(df: pd.DataFrame, audience_name: str = "") -> bytes:
    output = BytesIO()
    excel_df = df[STANDARD_COLUMNS].copy().where(pd.notna(df[STANDARD_COLUMNS]), "")
    max_row = 2000
    total_formula = f'=COUNTA(Profielen!$A$2:$A${max_row})'
    top_companies = excel_df["huidig_bedrijf"].replace("", "Onbekend").fillna("Onbekend").value_counts().head(20).index.tolist()
    top_sectors = excel_df["sector"].replace("", "Onbekend").fillna("Onbekend").value_counts().head(20).index.tolist()
    top_locations = excel_df["locatie"].replace("", "Onbekend").fillna("Onbekend").value_counts().head(20).index.tolist()
    top_education = EDUCATION_ORDER
    tenure_buckets = ["<1 jaar", "1-2 jaar", "3-5 jaar", "6-10 jaar", "10+ jaar", "Onbekend"]

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        excel_df.to_excel(writer, sheet_name="Profielen", index=False)
        workbook = writer.book
        workbook.set_calc_mode("auto")
        ws = writer.sheets["Profielen"]
        header_fmt = workbook.add_format({"bold": True, "font_color": "white", "bg_color": "#155E75", "border": 1})
        wrap = workbook.add_format({"text_wrap": True, "valign": "top"})
        for col_num, col in enumerate(excel_df.columns):
            ws.write(0, col_num, col, header_fmt)
            width = min(max(14, int(excel_df[col].astype(str).str.len().quantile(0.9) if len(excel_df) else 14)), 42)
            ws.set_column(col_num, col_num, width, wrap)
        ws.set_column(5, 5, 16)
        ws.freeze_panes(1, 1)
        ws.autofilter(0, 0, len(excel_df), len(excel_df.columns) - 1)

        summary = workbook.add_worksheet("Samenvatting")
        title_fmt = workbook.add_format({"bold": True, "font_color": "white", "bg_color": "#155E75", "font_size": 14})
        sub_fmt = workbook.add_format({"bold": True, "bg_color": "#E0F2FE"})
        header_fmt_2 = workbook.add_format({"bold": True, "bg_color": "#E0F2FE", "border": 1})
        pct_fmt = workbook.add_format({"num_format": "0.0%"})
        num_fmt = workbook.add_format({"num_format": "0.0"})
        summary.write(0, 0, "Market mapping - samenvatting", title_fmt)
        summary.write(1, 0, "Deze tab rekent mee wanneer je waarden op tab Profielen aanpast.", None)

        summary.write_row(3, 0, ["Metriek", "Waarde", "Percentage", "Opmerking"], header_fmt_2)
        kpis = [
            ("Profielen", total_formula, True, "Telt namen op tab Profielen"),
            ("Met huidig bedrijf", f'=COUNTIF(Profielen!$B$2:$B${max_row},"<>")', True, ""),
            ("Met tenure", f'=COUNT(Profielen!$F$2:$F${max_row})', True, ""),
            ("Gemiddelde tenure", f'=IFERROR(AVERAGE(Profielen!$F$2:$F${max_row}),"")', False, "In jaren"),
            ("Met eerdere rollen", f'=COUNTIF(Profielen!$G$2:$G${max_row},"<>")', True, ""),
            ("Met opleiding", f'=COUNTIFS(Profielen!$H$2:$H${max_row},"<>",Profielen!$H$2:$H${max_row},"<>Overig")', True, ""),
            ("Met activiteit", f'=COUNTIF(Profielen!$J$2:$J${max_row},"<>")', True, ""),
        ]
        for offset, (label, value_formula, show_pct, note) in enumerate(kpis, start=4):
            excel_row = offset + 1
            summary.write(offset, 0, label)
            summary.write_formula(offset, 1, value_formula, num_fmt if label == "Gemiddelde tenure" else None)
            if show_pct:
                pct_formula = '=IF($B$5=0,"",1)' if label == "Profielen" else f'=IF($B$5=0,"",$B{excel_row}/$B$5)'
                summary.write_formula(offset, 2, pct_formula, pct_fmt)
            summary.write(offset, 3, note)

        def write_count_block(start_row: int, title: str, labels: Iterable[str], source_col: str, chart: bool = True) -> int:
            labels = list(labels)
            summary.write(start_row, 0, title, sub_fmt)
            summary.write_row(start_row + 1, 0, ["categorie", "aantal", "percentage"], header_fmt_2)
            for i, label in enumerate(labels, start=start_row + 2):
                summary.write(i, 0, label)
                if label == "Onbekend":
                    formula = f'=COUNTIFS(Profielen!$A$2:$A${max_row},"<>",Profielen!${source_col}$2:${source_col}${max_row},"")'
                else:
                    safe_label = str(label).replace('"', '""')
                    formula = f'=COUNTIF(Profielen!${source_col}$2:${source_col}${max_row},"{safe_label}")'
                summary.write_formula(i, 1, formula)
                summary.write_formula(i, 2, f'=IF($B$5=0,"",$B{i + 1}/$B$5)', pct_fmt)
            if chart and labels:
                chart_obj = workbook.add_chart({"type": "column"})
                chart_end = start_row + 1 + min(len(labels), 10)
                chart_obj.add_series(
                    {
                        "name": title,
                        "categories": ["Samenvatting", start_row + 2, 0, chart_end, 0],
                        "values": ["Samenvatting", start_row + 2, 1, chart_end, 1],
                    }
                )
                chart_obj.set_title({"name": title})
                chart_obj.set_legend({"none": True})
                summary.insert_chart(start_row, 5, chart_obj, {"x_scale": 1.25, "y_scale": 1.1})
            return start_row + len(labels) + 4

        row = 13
        row = write_count_block(row, "Bedrijven", top_companies, "B", chart=False)
        row = write_count_block(row, "Sectoren", top_sectors, "E", chart=True)
        row = write_count_block(row, "Locaties", top_locations, "D", chart=True)
        row = write_count_block(row, "Opleiding", top_education, "H", chart=False)

        summary.write(row, 0, "Tenure-buckets", sub_fmt)
        summary.write_row(row + 1, 0, ["categorie", "aantal", "percentage"], header_fmt_2)
        tenure_formulas = [
            f'=COUNTIFS(Profielen!$F$2:$F${max_row},">=0",Profielen!$F$2:$F${max_row},"<1")',
            f'=COUNTIFS(Profielen!$F$2:$F${max_row},">=1",Profielen!$F$2:$F${max_row},"<=2")',
            f'=COUNTIFS(Profielen!$F$2:$F${max_row},">=3",Profielen!$F$2:$F${max_row},"<=5")',
            f'=COUNTIFS(Profielen!$F$2:$F${max_row},">=6",Profielen!$F$2:$F${max_row},"<=10")',
            f'=COUNTIF(Profielen!$F$2:$F${max_row},">10")',
            f'=$B$5-COUNT(Profielen!$F$2:$F${max_row})',
        ]
        for i, (label, formula) in enumerate(zip(tenure_buckets, tenure_formulas), start=row + 2):
            summary.write(i, 0, label)
            summary.write_formula(i, 1, formula)
            summary.write_formula(i, 2, f'=IF($B$5=0,"",$B{i + 1}/$B$5)', pct_fmt)

        summary.set_column(0, 0, 34)
        summary.set_column(1, 2, 14)
        summary.set_column(3, 3, 48)

        onepager = workbook.add_worksheet("One-pager")
        onepager.hide_gridlines(2)
        onepager.set_landscape()
        onepager.set_paper(9)
        onepager.fit_to_pages(1, 1)
        onepager.set_margins(left=0.25, right=0.25, top=0.25, bottom=0.25)
        for col in range(10):
            onepager.set_column(col, col, 13)
        for row_idx in range(38):
            onepager.set_row(row_idx, 24)

        dark_blue = "#073763"
        mid_blue = "#155E75"
        light_blue = "#EAF4FA"
        green = "#92D050"
        text_dark = "#1F2937"

        one_title_fmt = workbook.add_format(
            {"bold": True, "font_size": 20, "font_color": dark_blue, "valign": "vcenter"}
        )
        one_subtitle_fmt = workbook.add_format({"font_size": 10, "font_color": "#4B5563"})
        card_label_fmt = workbook.add_format(
            {"bold": True, "font_color": "white", "bg_color": mid_blue, "align": "center", "valign": "vcenter"}
        )
        card_value_fmt = workbook.add_format(
            {"bold": True, "font_size": 18, "font_color": dark_blue, "bg_color": light_blue, "align": "center", "valign": "vcenter", "border": 1, "border_color": "#B7C9D6"}
        )
        section_fmt = workbook.add_format(
            {"bold": True, "font_color": "white", "bg_color": dark_blue, "align": "left", "valign": "vcenter"}
        )
        table_header_fmt = workbook.add_format(
            {"bold": True, "font_color": text_dark, "bg_color": "#DDEBF7", "border": 1}
        )
        table_cell_fmt = workbook.add_format({"border": 1, "border_color": "#D9E2EA"})
        table_pct_fmt = workbook.add_format({"num_format": "0.0%", "border": 1, "border_color": "#D9E2EA"})
        note_fmt = workbook.add_format(
            {"text_wrap": True, "valign": "top", "border": 2, "border_color": green, "bg_color": "#F7FBF2", "font_color": text_dark}
        )
        green_bar_fmt = workbook.add_format({"bg_color": green})

        if HEADER_IMAGE_PATH.exists():
            onepager.insert_image("A1", str(HEADER_IMAGE_PATH), {"x_scale": 0.475, "y_scale": 0.475})
            onepager.merge_range("A10:J10", "", green_bar_fmt)
            content_start = 11
        else:
            onepager.merge_range("A1:J8", "Derks & Derks", workbook.add_format({"bold": True, "font_color": "white", "bg_color": dark_blue, "font_size": 22, "align": "center", "valign": "vcenter"}))
            onepager.merge_range("A9:J9", "", green_bar_fmt)
            content_start = 10

        page_title = f"Market Map {clean(audience_name) or 'Profielen'}"
        onepager.merge_range(content_start, 0, content_start, 6, page_title, one_title_fmt)
        onepager.merge_range(content_start + 1, 0, content_start + 1, 6, "Kernbeeld van de profielen in deze mapping", one_subtitle_fmt)

        table_start = content_start + 8
        total_expr = total_formula.lstrip("=")
        company_first_row = table_start + 3
        company_last_row = company_first_row + max(0, len(top_companies[:5]) - 1)
        sector_first_row = table_start + 12
        sector_last_row = sector_first_row + max(0, len(top_sectors[:5]) - 1)
        cards = [
            ("Profielen", total_formula),
            ("Gem. tenure", f'=IFERROR(AVERAGE(Profielen!$F$2:$F${max_row}),"")'),
            ("Topbedrijf", f'=IFERROR(INDEX($D${company_first_row}:$D${company_last_row},MATCH(MAX($E${company_first_row}:$E${company_last_row}),$E${company_first_row}:$E${company_last_row},0)),"")'),
            ("Topsector", f'=IFERROR(INDEX($D${sector_first_row}:$D${sector_last_row},MATCH(MAX($E${sector_first_row}:$E${sector_last_row}),$E${sector_first_row}:$E${sector_last_row},0)),"")'),
        ]
        card_cols = [0, 2, 4, 6]
        for (label, formula), col in zip(cards, card_cols):
            onepager.merge_range(content_start + 3, col, content_start + 3, col + 1, label, card_label_fmt)
            onepager.merge_range(content_start + 4, col, content_start + 5, col + 1, "", card_value_fmt)
            onepager.write_formula(content_start + 4, col, formula, card_value_fmt)

        onepager.merge_range(table_start, 0, table_start, 2, "Opleiding", section_fmt)
        onepager.write_row(table_start + 1, 0, ["categorie", "aantal", "%"], table_header_fmt)
        for i, label in enumerate(EDUCATION_ORDER, start=table_start + 2):
            onepager.write(i, 0, label, table_cell_fmt)
            safe_label = str(label).replace('"', '""')
            onepager.write_formula(i, 1, f'=COUNTIF(Profielen!$H$2:$H${max_row},"{safe_label}")', table_cell_fmt)
            onepager.write_formula(i, 2, f'=IF({total_expr}=0,"",$B{i + 1}/({total_expr}))', table_pct_fmt)

        onepager.merge_range(table_start, 3, table_start, 5, "Top bedrijven", section_fmt)
        onepager.write_row(table_start + 1, 3, ["bedrijf", "aantal", "%"], table_header_fmt)
        for i, label in enumerate(top_companies[:5], start=table_start + 2):
            onepager.write(i, 3, label, table_cell_fmt)
            safe_label = str(label).replace('"', '""')
            onepager.write_formula(i, 4, f'=COUNTIF(Profielen!$B$2:$B${max_row},"{safe_label}")', table_cell_fmt)
            onepager.write_formula(i, 5, f'=IF({total_expr}=0,"",$E{i + 1}/({total_expr}))', table_pct_fmt)

        onepager.merge_range(table_start + 9, 0, table_start + 9, 2, "Tenure", section_fmt)
        onepager.write_row(table_start + 10, 0, ["categorie", "aantal", "%"], table_header_fmt)
        for i, label in enumerate(tenure_buckets, start=table_start + 11):
            onepager.write(i, 0, label, table_cell_fmt)
            formula = tenure_formulas[i - (table_start + 11)]
            onepager.write_formula(i, 1, formula, table_cell_fmt)
            onepager.write_formula(i, 2, f'=IF({total_expr}=0,"",$B{i + 1}/({total_expr}))', table_pct_fmt)

        onepager.merge_range(table_start + 9, 3, table_start + 9, 5, "Top sectoren", section_fmt)
        onepager.write_row(table_start + 10, 3, ["sector", "aantal", "%"], table_header_fmt)
        for i, label in enumerate(top_sectors[:5], start=table_start + 11):
            onepager.write(i, 3, label, table_cell_fmt)
            safe_label = str(label).replace('"', '""')
            onepager.write_formula(i, 4, f'=COUNTIF(Profielen!$E$2:$E${max_row},"{safe_label}")', table_cell_fmt)
            onepager.write_formula(i, 5, f'=IF({total_expr}=0,"",$E{i + 1}/({total_expr}))', table_pct_fmt)

        onepager.merge_range(table_start, 7, table_start, 9, "Duiding / klantnotitie", section_fmt)
        onepager.merge_range(table_start + 1, 7, table_start + 17, 9, "Vul hier zelf de belangrijkste interpretatie, nuance of klantboodschap in.", note_fmt)
    return output.getvalue()


def main() -> None:
    st.set_page_config(page_title="Market Mapping", layout="wide")
    st.title("Market Mapping")
    st.caption("Upload een LinkedIn Recruiter Word-export of een Excel/CSV en maak een nette profielmapping met twijfelcheck.")

    with st.sidebar:
        st.header("Instellingen")
        sector_preset = "Automatisch"
        extra_exclude_terms = tokens(
            st.text_area(
                "Extra twijfel-/uitsluitingswoorden",
                "",
                help="Optioneel. De app gebruikt al een standaardlijst. Vul hier alleen extra woorden in die voor jouw zoekopdracht verdacht zijn.",
            )
        )

    uploaded = st.file_uploader("Upload Word, Excel of CSV", type=["docx", "xlsx", "xls", "csv"])

    if not uploaded:
        st.info("Upload een bestand om te beginnen.")
        return

    df = read_table(uploaded)
    if df.empty:
        st.error("Ik kon geen profielblokken herkennen. Probeer een Excel/CSV of controleer of de Word-export LinkedIn-profielregels bevat.")
        st.stop()

    sector_buckets, sector_rules, resolved_preset = sector_config(sector_preset, df)
    if sector_buckets:
        df["sector"] = df.apply(lambda row: classify_sector(row, sector_buckets, sector_rules), axis=1)
        st.caption(f"Gebruikte sectorindeling: {resolved_preset}")
    include_terms = DEFAULT_INCLUDE_TERMS.get(resolved_preset, DEFAULT_INCLUDE_TERMS["Pharma / biotech"])
    exclude_terms = DEFAULT_EXCLUDE_TERMS + extra_exclude_terms

    st.subheader("1. Ingelezen profielen")
    st.write(f"{len(df)} profielen ingelezen.")
    edited_df = st.data_editor(df, use_container_width=True, num_rows="dynamic", height=420)

    st.subheader("2. Twijfelprofielen")
    suspects = flag_outliers(edited_df, include_terms, exclude_terms)
    if suspects.empty:
        st.success("Geen duidelijke twijfelprofielen gevonden met de huidige keywords.")
        final_df = edited_df.copy()
    else:
        suspects = suspects.sort_values(["twijfelscore", "naam"], ascending=[False, True]).reset_index(drop=True)
        st.write("Vink profielen aan die uit de export moeten. Open hieronder een profiel voor extra context.")
        review_columns = [
            "verwijderen",
            "twijfelscore",
            "index",
            "naam",
            "huidig_bedrijf",
            "functietitel",
            "sector",
            "reden",
        ]
        edited_suspects = st.data_editor(
            suspects[review_columns],
            use_container_width=True,
            hide_index=True,
            height=280,
            column_config={
                "verwijderen": st.column_config.CheckboxColumn("verwijderen"),
                "twijfelscore": st.column_config.NumberColumn("score", width="small"),
                "index": st.column_config.NumberColumn("rij", width="small"),
                "reden": st.column_config.TextColumn("reden", width="large"),
            },
        )
        st.caption("Profielcontext voor beoordeling")
        for _, suspect in suspects.iterrows():
            title = f"{suspect['naam']} - {suspect['functietitel'] or 'functie onbekend'}"
            with st.expander(title):
                st.write(f"**Bedrijf:** {suspect['huidig_bedrijf'] or '-'}")
                st.write(f"**Sector:** {suspect['sector'] or '-'}")
                st.write(f"**Locatie:** {suspect['locatie'] or '-'}")
                st.write(f"**Tenure huidige rol:** {suspect['tenure_huidige_rol_jaren'] or '-'}")
                st.write(f"**Eerdere rollen:** {suspect['info_eerdere_rollen'] or '-'}")
                st.write(f"**Opleiding:** {suspect['hoogst_afgeronde_opleiding'] or '-'}")
                st.write(f"**Interesse signalen:** {suspect['interesse_signalen'] or '-'}")
                st.write(f"**Activiteit:** {suspect['activiteit'] or '-'}")
                st.write(f"**Waarom twijfel:** {suspect['reden'] or '-'}")
        remove_idx = set(edited_suspects.loc[edited_suspects["verwijderen"], "index"].tolist())
        final_df = edited_df.drop(index=remove_idx).reset_index(drop=True)
        st.info(f"{len(remove_idx)} profiel(en) gemarkeerd voor verwijdering. Export bevat straks {len(final_df)} profielen.")

    st.subheader("3. Export")
    suggested_audience = guess_audience_name(final_df)
    audience_name = st.text_input("Naam doelgroep voor bestandsnaam", suggested_audience)
    st.caption("De download bevat tab 1 Profielen, tab 2 Samenvatting en tab 3 One-pager.")
    excel = make_excel(final_df, audience_name)
    pdf = make_onepager_pdf(final_df, audience_name)
    col_excel, col_pdf = st.columns(2)
    with col_excel:
        st.download_button(
            "Download market mapping Excel",
            data=excel,
            file_name=safe_excel_filename(audience_name),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    with col_pdf:
        st.download_button(
            "Download klant-onepager PDF",
            data=pdf,
            file_name=safe_pdf_filename(audience_name),
            mime="application/pdf",
        )


if __name__ == "__main__":
    main()
